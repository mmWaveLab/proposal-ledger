#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from io import BytesIO
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ModuleNotFoundError as exc:
    raise SystemExit("缺少 python-docx，请先运行: python3 -m pip install -r requirements.txt") from exc


ROOT = Path(__file__).resolve().parents[1]
APPLICATIONS_DIR = ROOT / "applications"
SOURCE_NAME = "申报书.md"
REQUIRED_HEADINGS = [
    "一、需求分析",
    "二、目的用途",
    "三、实现目标",
    "经费数量",
    "四、取得效果或结果",
    "五、其他",
]


def set_font(run, size: float = 12, bold: bool = False, east: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.bold = bold


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    heading = doc.styles["Heading 1"]
    heading.font.name = "Times New Roman"
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading.font.size = Pt(14)
    heading.font.bold = True
    return doc


def add_para(
    doc: Document,
    text: str = "",
    *,
    style: str | None = None,
    first_line: bool = True,
    size: float = 12,
    bold: bool = False,
    east: str = "宋体",
    align: WD_ALIGN_PARAGRAPH | None = None,
):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        set_font(run, size=size, bold=bold, east=east)
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(6)
    if first_line and style is None:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    if align:
        paragraph.alignment = align
    return paragraph


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Heading 1")
    run = paragraph.add_run(text)
    set_font(run, size=14, bold=True, east="黑体")
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_font(run)
    paragraph.paragraph_format.line_spacing = 1.3
    paragraph.paragraph_format.space_after = Pt(4)


def set_cell_text(cell, text: str, *, bold: bool = False, align=None, size: float = 10.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(rows):
        for col_index in range(width):
            text = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            is_header = row_index == 0
            set_cell_text(
                cell,
                text,
                bold=is_header,
                align=WD_ALIGN_PARAGRAPH.CENTER if is_header or col_index in {1, 2, 3} else None,
                size=10 if col_index == width - 1 else 10.5,
            )
            if is_header:
                shade_cell(cell, "D9EAF7")
    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip()
        cells = [cell.strip() for cell in raw.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_image(doc: Document, source: Path, line: str) -> None:
    match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line.strip())
    if not match:
        return
    caption, target = match.groups()
    image = (source.parent / target).resolve()
    if caption:
        add_para(doc, caption, first_line=False, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    if image.exists():
        if image.suffix.lower() == ".webp":
            try:
                from PIL import Image
            except ModuleNotFoundError as exc:
                raise SystemExit("嵌入 WebP 图片需要 Pillow，请先运行: python3 -m pip install -r requirements.txt") from exc
            buffer = BytesIO()
            with Image.open(image) as opened:
                opened.convert("RGB").save(buffer, format="PNG")
            buffer.seek(0)
            doc.add_picture(buffer, width=Cm(13.5))
        else:
            doc.add_picture(str(image), width=Cm(13.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"图片文件缺失：{target}", first_line=False, size=10.5)


def generate(source: Path) -> Path:
    app_dir = source.parent
    output = app_dir / f"{app_dir.name}-申报书.docx"
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = setup_document()

    index = 0
    if lines and lines[0].startswith("# "):
        title = lines[0][2:].strip()
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title)
        set_font(run, size=18, bold=True, east="黑体")
        paragraph.paragraph_format.space_after = Pt(14)
        index = 1

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip())
        elif stripped.startswith("- "):
            add_bullet(doc, stripped[2:].strip())
        elif stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        elif stripped.startswith("!["):
            add_image(doc, source, stripped)
        else:
            add_para(doc, stripped)
        index += 1

    doc.save(output)
    return output


def find_sources(paths: list[str], all_sources: bool) -> list[Path]:
    if all_sources:
        return sorted(APPLICATIONS_DIR.glob(f"**/{SOURCE_NAME}"))
    if not paths:
        raise SystemExit("请指定申报目录或使用 --all")
    sources: list[Path] = []
    for raw in paths:
        path = Path(raw)
        if not path.is_absolute():
            path = ROOT / path
        sources.append(path if path.name == SOURCE_NAME else path / SOURCE_NAME)
    return sources


def validate_source(source: Path) -> list[str]:
    errors: list[str] = []
    if not source.exists():
        return [f"缺少源文件: {source.relative_to(ROOT)}"]
    text = source.read_text(encoding="utf-8")
    headings = re.findall(r"^##\s+(.+?)\s*$", text, re.MULTILINE)
    if headings != REQUIRED_HEADINGS:
        errors.append(
            f"{source.relative_to(ROOT)}: 二级标题必须严格为 "
            + "、".join(REQUIRED_HEADINGS)
        )
    return errors


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate proposal DOCX files from 申报书.md")
    parser.add_argument("paths", nargs="*", help="application directories or 申报书.md files")
    parser.add_argument("--all", action="store_true", help="generate every applications/**/申报书.md")
    parser.add_argument("--check-source", action="store_true", help="only validate source structure")
    args = parser.parse_args()

    sources = find_sources(args.paths, args.all)
    errors: list[str] = []
    for source in sources:
        errors.extend(validate_source(source))
    if errors:
        print("Proposal source validation failed:")
        for error in errors:
            print(f"- {error}")
        return 1
    if args.check_source:
        print(f"Proposal source validation passed: {len(sources)} source file(s)")
        return 0

    for source in sources:
        output = generate(source)
        print(output.relative_to(ROOT))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
